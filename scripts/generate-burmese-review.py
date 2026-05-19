#!/usr/bin/env python3
"""Generate a Burmese-only Word document for uncle to review every Burmese string.

Reads messages/my.json, writes:
  - ~/Downloads/swh-burmese-review.docx
  - ~/Downloads/swh-burmese-review-keymap.json

The keymap maps entry labels (e.g. "1", "5a") back to JSON key paths so uncle's
edits can be round-tripped into messages/my.json later.

Requires: python-docx  (pip install python-docx)
"""
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parent.parent
MY_JSON = REPO_ROOT / "messages" / "my.json"
OUTPUT_DOCX = Path.home() / "Downloads" / "swh-burmese-review.docx"
OUTPUT_KEYMAP = Path.home() / "Downloads" / "swh-burmese-review-keymap.json"

BURMESE_FONT = "Noto Sans Myanmar"
BODY_PT = 14
HEADING_PT = 20
TITLE_PT = 22
REF_PT = 7
REF_COLOR = RGBColor(0x88, 0x88, 0x88)
RED = RGBColor(0xC0, 0x00, 0x00)

EMPTY_MARKER = "[မရှိသေး — ထည့်ပါ]"
PLACEHOLDER_NOTE = (
    "{ } အတွင်းရှိ စာသားကို မပြောင်းပါနှင့်။ "
    "၎င်းသည် အလိုအလျောက် ဖြည့်စွက်သော ကိန်းသေဖြစ်သည်။"
)
PLACEHOLDER_RE = re.compile(r"\{[^}]+\}")

# Burmese labels for top-level JSON sections (pulled from Navigation.links.* where possible)
SECTION_TITLES = {
    "Metadata": "မီတာဒေတာ",
    "Navigation": "ဝက်ဘ်ဆိုက် လမ်းညွှန်",
    "Footer": "အောက်ခြေ",
    "Common": "ဘုံ",
    "home": "ပင်မ",
    "about": "ကျွန်ုပ်တို့အကြောင်း",
    "fisheries": "ရေလုပ်ငန်း",
    "distribution": "ကုန်သွယ်ရေးနှင့် ဖြန့်ဖြူးရေး",
    "fertilizer": "ဓါတ်မြေသြဇာ",
    "csr": "CSR",
    "contact": "ဆက်သွယ်ရန်",
}

# AI-placeholder Burmese strings flagged across phase docs (E6, E10, E17/E17.1/E17.2)
# as pending uncle review. Highlighted in yellow.
PRIORITY_KEYS = {
    # E6 — Skechers block + home.stats
    "home.stats.eyebrow",
    "home.stats.heading",
    "distribution.skechers.eyebrow",
    "distribution.skechers.heading",
    "distribution.skechers.body",
    "distribution.skechers.bullets.item1",
    "distribution.skechers.bullets.item2",
    "distribution.skechers.bullets.item3",
    # E10 — contact form illustrative note
    "contact.form.illustrativeNote",
    # E17 / E17.1 / E17.2 — MicroLife / SWH parent eyebrow / OWNED BRAND body
    "fertilizer.hero.eyebrow",
    "fertilizer.hero.subtitle",
    "fertilizer.overview.body",
    "fertilizer.overview.bullets.item2",
    "fertilizer.overview.bullets.item3",
    "fertilizer.overview.bullets.item4",
    "fertilizer.overview.bullets.item5",
}


def apply_burmese_font(run, size_pt):
    """Force a run to use Noto Sans Myanmar at the given size for all script slots."""
    run.font.name = BURMESE_FONT
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), BURMESE_FONT)
    # Complex-script size (Burmese is a complex script in Word)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(szCs)


def collect_entries_for_section(section_data, root_key):
    """Walk a top-level section and return logical entries in JSON order.

    Each entry is a dict with:
      key:                full dotted JSON path
      value:              string value
      is_array_item:      True if originated from a JSON array
      array_parent_key:   dotted path of the parent array (or None)
    """
    entries = []

    def walk(node, path):
        if isinstance(node, dict):
            for k, v in node.items():
                walk(v, path + [k])
        elif isinstance(node, list):
            parent_key = ".".join(path)
            for i, v in enumerate(node):
                if isinstance(v, str):
                    entries.append({
                        "key": ".".join(path + [str(i)]),
                        "value": v,
                        "is_array_item": True,
                        "array_parent_key": parent_key,
                    })
                elif isinstance(v, (dict, list)):
                    walk(v, path + [str(i)])
        elif isinstance(node, str):
            entries.append({
                "key": ".".join(path),
                "value": node,
                "is_array_item": False,
                "array_parent_key": None,
            })
        # ignore numbers/booleans/null

    walk(section_data, [root_key])
    return entries


def add_burmese_paragraph(doc, text, size_pt=BODY_PT, bold=False, color=None,
                          highlight=None, bullet=False):
    p = doc.add_paragraph(style="List Bullet") if bullet else doc.add_paragraph()
    run = p.add_run(text)
    apply_burmese_font(run, size_pt)
    if bold:
        run.bold = True
    if color is not None:
        run.font.color.rgb = color
    if highlight is not None:
        run.font.highlight_color = highlight
    return p, run


def main():
    data = json.loads(MY_JSON.read_text(encoding="utf-8"))

    doc = Document()
    # Default style fallback — set Normal to Burmese font
    normal = doc.styles["Normal"]
    normal.font.name = BURMESE_FONT
    normal.font.size = Pt(BODY_PT)

    # --- COVER PAGE ---
    add_burmese_paragraph(
        doc, "SWH ဝက်ဘ်ဆိုက် မြန်မာစာ ပြန်လည်စိစစ်ခြင်း",
        size_pt=TITLE_PT, bold=True,
    )
    add_burmese_paragraph(doc, "၂၀၂၆ ခုနှစ်၊ မေလ ၁၉ ရက်", size_pt=14)
    doc.add_paragraph()

    for instr in [
        "မြန်မာ စာသားများကို တိုက်ရိုက် ပြုပြင်ပါ",
        "ပြီးဆုံးပါက ဤဖိုင်ကို ပြန်ပို့ပေးပါ",
        "နံပါတ်များနှင့် မီးခိုးရောင် ကုဒ်များကို မပြောင်းပါနှင့်",
    ]:
        add_burmese_paragraph(doc, instr, size_pt=14, bullet=True)

    # --- SECTIONS ---
    keymap = {}
    entry_counter = 0
    placeholder_note_shown = False
    uncategorized = []

    for section_key, section_data in data.items():
        title = SECTION_TITLES.get(section_key)
        if title is None:
            uncategorized.append(section_key)
            title = section_key

        # Page break before each section
        pb_p = doc.add_paragraph()
        pb_p.add_run().add_break(WD_BREAK.PAGE)

        heading_p = doc.add_heading(level=1)
        heading_run = heading_p.add_run(title)
        apply_burmese_font(heading_run, HEADING_PT)

        entries = collect_entries_for_section(section_data, section_key)

        current_array_parent = None
        array_letter_offset = 0

        for ent in entries:
            if ent["is_array_item"]:
                if ent["array_parent_key"] != current_array_parent:
                    entry_counter += 1
                    current_array_parent = ent["array_parent_key"]
                    array_letter_offset = 0
                letter = chr(ord("a") + array_letter_offset)
                array_letter_offset += 1
                label = f"{entry_counter}{letter}"
            else:
                current_array_parent = None
                entry_counter += 1
                label = str(entry_counter)

            keymap[label] = ent["key"]

            # Entry line: number + Burmese text
            text_p = doc.add_paragraph()
            num_run = text_p.add_run(f"{label}. ")
            apply_burmese_font(num_run, BODY_PT)
            num_run.bold = True

            value = ent["value"]
            is_priority = ent["key"] in PRIORITY_KEYS

            if value == "":
                empty_run = text_p.add_run(EMPTY_MARKER)
                apply_burmese_font(empty_run, BODY_PT)
                empty_run.font.color.rgb = RED
            else:
                val_run = text_p.add_run(value)
                apply_burmese_font(val_run, BODY_PT)
                if is_priority:
                    val_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

            # Reference code (gray, tiny)
            ref_p = doc.add_paragraph()
            ref_run = ref_p.add_run(f"[ကုဒ်: {ent['key']}]")
            apply_burmese_font(ref_run, REF_PT)
            ref_run.font.color.rgb = REF_COLOR

            # Placeholder note on first occurrence anywhere in the doc
            if (not placeholder_note_shown
                    and value
                    and PLACEHOLDER_RE.search(value)):
                note_p = doc.add_paragraph()
                note_run = note_p.add_run(PLACEHOLDER_NOTE)
                apply_burmese_font(note_run, REF_PT)
                note_run.font.color.rgb = REF_COLOR
                placeholder_note_shown = True

            doc.add_paragraph()  # blank line between entries

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    OUTPUT_KEYMAP.write_text(
        json.dumps(keymap, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    total = len(keymap)
    priority_in_doc = sum(1 for k in keymap.values() if k in PRIORITY_KEYS)
    priority_not_found = sorted(PRIORITY_KEYS - set(keymap.values()))

    print(f"Wrote {OUTPUT_DOCX}")
    print(f"Wrote {OUTPUT_KEYMAP}")
    print(f"Total entries:               {total}")
    print(f"Priority entries highlighted: {priority_in_doc}")
    if priority_not_found:
        print(f"Priority keys NOT found in JSON: {priority_not_found}")
    if uncategorized:
        print(f"Uncategorized top-level keys:    {uncategorized}")
    else:
        print("All top-level keys categorized under known page sections.")


if __name__ == "__main__":
    main()
